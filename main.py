import fitz  
import os, time, sys
from datetime import date
import pathlib
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches

## FUNÇÕES =============================================================
def pdf_para_imagens(pdf_path, output_folder):
    doc = fitz.open(pdf_path)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    for i in range(len(doc)):
        pagina = doc.load_page(i)
        imagem = pagina.get_pixmap(dpi=300)
        imagem_path = os.path.join(output_folder, f"pagina_{i + 1}.png")
        imagem.save(imagem_path)

def retornarData():
    data = str(date.today()).split("-")
    datacorreta = f"{data[2]}/{data[1]}/{data[0]}"
    return datacorreta

def formatarData(celula):
    celula.text = retornarData()
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Aharoni"
            run.font.size = Pt(36)
    cell_xml_element = celula._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#FFF2CC")
    table_cell_properties.append(shade_obj)    

## =====================================================================
print(r"""
________            _____________                           __________  
___  __ \_________________  /__(_)______ _______________    ___  ___  \ 
__  /_/ /_  ___/  _ \  __  /__  /__  __ `__ \  _ \  ___/    __  / _ \  |
_  ____/_  /   /  __/ /_/ / _  / _  / / / / /  __/ /__      _  / , _/ / 
/_/     /_/    \___/\__,_/  /_/  /_/ /_/ /_/\___/\___/      | /_/|_| /  
                                                             \______/   
Iniciando processo de formatação...
      """)

pdf_relatorio_oleo = '' 
documento_word_nome = ''


try: 
    arquivos = os.listdir('./')
    for arquivo in arquivos:
        if arquivo.endswith(".pdf"):
            pdf_relatorio_oleo = arquivo
        if arquivo.endswith(".docx"):
            documento_word_nome = arquivo
    w = open(documento_word_nome, 'rb')
    documentoWord = Document(w)
except:
    print("ERRO: Erro ao identificar arquivos para formatação.")
    time.sleep(10)
    sys.exit(1)
    
caminhoWord = os.path.join("RELATÓRIO FORMATADO", documento_word_nome)
os.makedirs("RELATÓRIO FORMATADO", exist_ok=True)
pdf_para_imagens(pdf_relatorio_oleo, "imagens_pdf")
formatarData(documentoWord.tables[0].columns[0].cells[0])

teste = 1
nmImagens = len(os.listdir("./imagens_pdf"))

for i in documentoWord.paragraphs:
    if i.text == "[LISTA-IMAGENS]":
        for j in os.listdir("./imagens_pdf"):
            if teste != nmImagens:
                paragrafo = documentoWord.add_paragraph()
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = paragrafo.add_run()
                img.add_picture(rf"{str(pathlib.Path().resolve())}\imagens_pdf\{j}", width=Inches(6))
                documentoWord.add_page_break() 
            else:
                paragrafo = documentoWord.add_paragraph()
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = paragrafo.add_run()
                img.add_picture(rf"{str(pathlib.Path().resolve())}\imagens_pdf\{j}", width=Inches(6))
            teste +=1
        p = i._element
        p.getparent().remove(p)
        p._p = p._element = None
documentoWord.save(caminhoWord)
shutil.rmtree('./imagens_pdf')

print("\nArquivos formatados com sucesso!\n")
time.sleep(10)